#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
JH Coaching-Protokoll Generator (Dual-Brand)
Unterstuetzt: Praxis Factory GmbH + JH Praxismanagement

Verwendung:
  python generate.py <input.json> <output.docx>

Das JSON muss ein Feld "brand" enthalten:
  "praxisfactory"    -> Praxis Factory GmbH Design
  "praxismanagement" -> JH Praxismanagement Design

Abhaengigkeiten:
  pip install python-docx
"""

import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# BRAND-DEFINITIONEN
# ============================================================

BRANDS = {
    "praxisfactory": {
        "primary_dark": RGBColor(0x1B, 0x3A, 0x5C),
        "primary": RGBColor(0x23, 0x4B, 0x6E),
        "light_bg": "EEF5D5",
        "header_line": "1B3A5C",
        "table_header_bg": "1B3A5C",
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "session_header_bg": "EEF5D5",
        "session_header_text_rgb": RGBColor(0x1B, 0x3A, 0x5C),
        "quote_bg": "EEF5D5",
        "table_border": "C2D530",
        "template": "template-praxisfactory.docx",
        "coach_name": "Juliane",
    },
    "praxismanagement": {
        "primary_dark": RGBColor(0x3D, 0x5A, 0x6B),
        "primary": RGBColor(0x7A, 0x9B, 0xA5),
        "light_bg": "E4EDF0",
        "header_line": "3D5A6B",
        "table_header_bg": "3D5A6B",
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "session_header_bg": "E4EDF0",
        "session_header_text_rgb": RGBColor(0x3D, 0x5A, 0x6B),
        "quote_bg": "E4EDF0",
        "table_border": "7A9BA5",
        "template": "template-praxismanagement.docx",
        "coach_name": "Juliane",
    }
}

FONT = "Arial"

# ============================================================
# STYLE-SETUP
# ============================================================

def setup_styles(doc, brand):
    styles = doc.styles
    B = BRANDS[brand]

    if 'DT Titel' not in [s.name for s in styles]:
        title_style = styles.add_style('DT Titel', WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = styles['DT Titel']
    title_style.font.name = FONT
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = B["primary_dark"]
    title_style.paragraph_format.space_after = Pt(6)

    if 'DT Untertitel' not in [s.name for s in styles]:
        subtitle_style = styles.add_style('DT Untertitel', WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle_style = styles['DT Untertitel']
    subtitle_style.font.name = FONT
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.color.rgb = B["primary"]
    subtitle_style.paragraph_format.space_after = Pt(12)

    if 'DT Meta' not in [s.name for s in styles]:
        meta_style = styles.add_style('DT Meta', WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta_style = styles['DT Meta']
    meta_style.font.name = FONT
    meta_style.font.size = Pt(10)
    meta_style.font.color.rgb = B["primary"]
    meta_style.paragraph_format.space_after = Pt(18)

    try:
        h1_style = styles['Heading 1']
    except KeyError:
        h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = FONT
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.color.rgb = B["primary_dark"]
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.keep_with_next = True

    try:
        h2_style = styles['Heading 2']
    except KeyError:
        h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = FONT
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.font.color.rgb = B["primary"]
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(4)
    h2_style.paragraph_format.keep_with_next = True

    try:
        body_style = styles['Normal']
    except KeyError:
        body_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = FONT
    body_style.font.size = Pt(10)
    body_style.paragraph_format.space_after = Pt(6)

    try:
        list_style = styles['List Bullet']
    except KeyError:
        list_style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    list_style.font.name = FONT
    list_style.font.size = Pt(10)

# ============================================================
# HILFSFUNKTIONEN
# ============================================================

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        mar = OxmlElement(f'w:{side}')
        mar.set(qn('w:w'), str(value))
        mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    tcPr.append(tcMar)

def set_table_borders(table, color):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBorders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBorders)

def set_table_width(table, pct=5000):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), str(pct))
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblPr.append(tblW)

def set_column_widths(table, widths_dxa):
    tbl = table._tbl
    tblGrid = tbl.tblGrid
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    else:
        for old in list(tblGrid):
            tblGrid.remove(old)
    for width in widths_dxa:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_dxa):
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(widths_dxa[idx]))
                tcW.set(qn('w:type'), 'dxa')
                for old in tcPr.findall(qn('w:tcW')):
                    tcPr.remove(old)
                tcPr.insert(0, tcW)

def prevent_row_break(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:cantSplit')):
        trPr.remove(old)
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)

def keep_row_with_next(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            pPr = para._p.get_or_add_pPr()
            for old in pPr.findall(qn('w:keepNext')):
                pPr.remove(old)
            keepNext = OxmlElement('w:keepNext')
            pPr.append(keepNext)

def prevent_table_split(table):
    rows = list(table.rows)
    for i, row in enumerate(rows):
        prevent_row_break(row)
        if i < len(rows) - 1:
            keep_row_with_next(row)

def set_keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:keepNext')):
        pPr.remove(old)
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)

def set_keep_together(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:keepLines')):
        pPr.remove(old)
    keepLines = OxmlElement('w:keepLines')
    pPr.append(keepLines)

def add_bottom_border(para, color, size=24):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def clear_body(doc):
    body = doc._body._body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

# ============================================================
# DOKUMENTELEMENTE
# ============================================================

def add_title_block(doc, kunde, datum, dauer, format_typ):
    doc.add_paragraph("Coaching-Protokoll", style='DT Titel')
    doc.add_paragraph(kunde, style='DT Untertitel')
    doc.add_paragraph(
        f"Datum: {datum}  |  Dauer: {dauer} Minuten  |  Format: {format_typ}",
        style='DT Meta'
    )

def add_h1(doc, num, text, brand):
    B = BRANDS[brand]
    p = doc.add_paragraph(f"{num}. {text}", style='Heading 1')
    add_bottom_border(p, B["header_line"], 8)
    set_keep_with_next(p)
    set_keep_together(p)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph(text, style='Heading 2')
    set_keep_with_next(p)
    return p

def add_text(doc, text):
    p = doc.add_paragraph(text, style='Normal')
    set_keep_together(p)
    return p

def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')

def add_quote_box(doc, text, brand):
    B = BRANDS[brand]
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    remove_table_borders(table)
    prevent_table_split(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, B["quote_bg"])
    set_cell_margins(cell, 120, 120, 150, 150)
    p = cell.paragraphs[0]
    p.style = 'Normal'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'"{text}"')
    run.italic = True
    doc.add_paragraph()

def add_massnahmen_table(doc, massnahmen, brand):
    B = BRANDS[brand]
    table = doc.add_table(rows=1 + len(massnahmen), cols=3)
    set_table_width(table)
    set_table_borders(table, B["table_border"])
    set_column_widths(table, [1400, 5600, 2000])
    prevent_table_split(table)

    headers = ['Wer', 'Was', 'Bis wann']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, B["table_header_bg"])
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = B["table_header_text"]
        run.font.name = FONT

    for row_idx, m in enumerate(massnahmen):
        row = table.rows[row_idx + 1]
        wer = m.get('wer', '')
        if wer.lower() == 'coach':
            wer = B["coach_name"]
        values = [wer, m.get('was', ''), m.get('bis_wann', '')]
        for col_idx, val in enumerate(values):
            cell = row.cells[col_idx]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = FONT
    doc.add_paragraph()

def add_sessions_table(doc, sessions, brand):
    B = BRANDS[brand]
    table = doc.add_table(rows=1 + len(sessions), cols=3)
    set_table_width(table)
    set_table_borders(table, B["table_border"])
    set_column_widths(table, [4000, 2000, 3000])
    prevent_table_split(table)

    headers = ['Session', 'Datum', 'Status']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, B["session_header_bg"])
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = B["session_header_text_rgb"]
        run.font.name = FONT

    for row_idx, s in enumerate(sessions):
        row = table.rows[row_idx + 1]
        values = [s.get('session', ''), s.get('datum', ''), s.get('status', '')]
        for col_idx, val in enumerate(values):
            cell = row.cells[col_idx]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = FONT

# ============================================================
# HAUPTFUNKTION
# ============================================================

def generate_protokoll(data, template_path, output_path, brand):
    doc = Document(template_path)
    clear_body(doc)
    setup_styles(doc, brand)

    add_title_block(
        doc,
        data.get('kunde', 'Kunde'),
        data.get('datum', ''),
        data.get('dauer', '60'),
        data.get('format', 'Zoom')
    )

    n = 0

    # 1. Anliegen & Ziel
    n += 1
    add_h1(doc, n, "Anliegen & Ziel der Sitzung", brand)
    if data.get('anliegen'):
        add_text(doc, data['anliegen'])

    # 2. Kernthemen
    n += 1
    add_h1(doc, n, "Kernthemen der aktuellen Sitzung", brand)
    for thema in data.get('kernthemen', []):
        add_h2(doc, thema.get('titel', 'Thema'))
        if thema.get('einleitung'):
            add_text(doc, thema['einleitung'])
        for punkt in thema.get('punkte', []):
            add_bullet(doc, punkt)
        if thema.get('erkenntnis'):
            p = doc.add_paragraph(style='Normal')
            p.add_run("\u2192 ").bold = False
            run = p.add_run(thema['erkenntnis'])
            run.italic = True

    # 3. Erkenntnisse & Aha-Momente
    n += 1
    add_h1(doc, n, "Erkenntnisse & Aha-Momente", brand)
    for e in data.get('erkenntnisse', []):
        add_bullet(doc, e)
    if data.get('zitat'):
        doc.add_paragraph()
        add_quote_box(doc, data['zitat'], brand)

    # 4. Reflexion & Feedback
    n += 1
    add_h1(doc, n, "Reflexion & Feedback", brand)
    for r in data.get('reflexion', []):
        add_text(doc, r)

    # 5. Vereinbarte Massnahmen
    n += 1
    add_h1(doc, n, "Vereinbarte Ma\u00dfnahmen", brand)
    if data.get('massnahmen'):
        add_massnahmen_table(doc, data['massnahmen'], brand)

    # 6. Sonstiges
    n += 1
    add_h1(doc, n, "Sonstiges", brand)
    sonstiges = data.get('sonstiges', [])
    if sonstiges:
        for item in sonstiges:
            add_bullet(doc, item)
    else:
        add_text(doc, "-")

    # 7. Status & naechste Termine
    n += 1
    add_h1(doc, n, "Status & n\u00e4chste Termine", brand)
    if data.get('sessions'):
        add_sessions_table(doc, data['sessions'], brand)

    doc.save(output_path)
    print(f"Protokoll erstellt: {output_path}")
    print(f"Brand: {brand} ({'Praxis Factory GmbH' if brand == 'praxisfactory' else 'JH Praxismanagement'})")

# ============================================================
# MAIN
# ============================================================

def main():
    if len(sys.argv) < 3:
        print("Verwendung: python generate.py <input.json> <output.docx>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    brand = data.get('brand', '').lower().strip()
    if brand not in BRANDS:
        print(f"FEHLER: Unbekannter Brand '{brand}'.")
        print(f"Gueltige Werte: {', '.join(BRANDS.keys())}")
        sys.exit(1)

    script_dir = Path(__file__).parent.parent
    template_filename = BRANDS[brand]["template"]
    template_path = script_dir / "assets" / template_filename

    if not template_path.exists():
        print(f"FEHLER: Template nicht gefunden: {template_path}")
        sys.exit(1)

    generate_protokoll(data, str(template_path), output_path, brand)

if __name__ == "__main__":
    main()
